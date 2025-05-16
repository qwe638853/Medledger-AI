import math
from tqdm import tqdm
import pymssql
from langchain_ollama import OllamaLLM
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.docstore.document import Document
from pydantic import BaseModel
from typing import Optional, Dict, Any, List, Tuple
from passlib.context import CryptContext
import PyPDF2
from docx import Document as DocxDocument
from io import BytesIO
from datetime import datetime
import logging
import json
import os
import shutil
from tenacity import retry, stop_after_attempt, wait_fixed
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
import torch
import time

# 載入 .env 檔案
load_dotenv()

# 初始化密碼加密工具
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# 設定日誌
logging.basicConfig(level=logging.INFO, format='%(levelname)s:%(name)s:%(message)s')
logger = logging.getLogger(__name__)

# 定義角色常量
ROLE_USER = "user"
ROLE_OTHER = "other"
ROLE_HEALTH_CENTER = "health_center"
VALID_ROLES = [ROLE_USER, ROLE_OTHER, ROLE_HEALTH_CENTER]

# 初始化嵌入模型（使用 GPU，全局範圍）
device = "cuda" if torch.cuda.is_available() else "cpu"
logger.info(f"使用設備: {device}")
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
st_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2", device=device)

# 初始化 Chroma 向量資料庫（全局變數）
persist_dir = "./chroma_db"
vectorstore = None
try:
    logger.info("初始化 Chroma 向量數據庫...")
    if os.path.exists(persist_dir):
        logger.info(f"使用現有的 Chroma 資料庫目錄 {persist_dir}")
    else:
        os.makedirs(persist_dir, exist_ok=True)
    vectorstore = Chroma(
        embedding_function=embedding_model,
        collection_name="health_knowledge",
        persist_directory=persist_dir
    )
    logger.info("Chroma 向量數據庫初始化成功")
except Exception as e:
    logger.error(f"初始化 Chroma 失敗: {str(e)}")
    raise Exception(f"初始化 Chroma 失敗: {str(e)}")

# 自定義表單類，用於簡化 OAuth2 登入表單
class CustomOAuth2PasswordRequestForm(BaseModel):
    username: str
    password: str
    scope: Optional[str] = None

    class Config:
        json_schema_extra = {
            "example": {
                "username": "A123456789",
                "password": "testpassword",
                "scope": "role:user"
            }
        }

# 定義請求體結構（用於註冊）
class RegisterRequest(BaseModel):
    full_name: str
    gender: str
    birth_date: str
    id_number: str
    password: str
    phone_number: str
    email: str
    role: str

# 定義請求體結構（用於登入）
class LoginRequest(BaseModel):
    id_number: str
    password: str
    role: str

# 定義請求體結構（用於忘記密碼）
class ForgotPasswordRequest(BaseModel):
    id_number: str
    email: str

# 定義請求體結構（用於互動模式）
class InteractiveRequest(BaseModel):
    query: str

# 連接到 Azure SQL Database 的通用函數
@retry(stop=stop_after_attempt(3), wait=wait_fixed(5))
def connect_to_db(db_config: Dict[str, Any]):
    try:
        conn = pymssql.connect(
            server=db_config['server'],
            port=db_config['port'],
            user=db_config['user'],
            password=db_config['password'],
            database=db_config['database'],
            login_timeout=30,
            charset='UTF-8'
        )
        logger.info("成功連接到 Azure SQL Database")
        return conn
    except Exception as e:
        logger.error(f"連線 Azure SQL Database 時發生錯誤: {str(e)}")
        raise Exception(f"連線 Azure SQL Database 時發生錯誤: {str(e)}")

# 檢查身分證字號是否已存在
def check_id_number_exists(db_config: Dict[str, Any], id_number: str) -> bool:
    try:
        conn = connect_to_db(db_config)
        cursor = conn.cursor()
        query = "SELECT COUNT(*) FROM users WHERE id_number = %s AND id_number != ''"
        cursor.execute(query, (id_number,))
        count = cursor.fetchone()[0]
        conn.close()
        logger.info(f"檢查身分證字號是否存在: id_number={id_number}, 存在={count > 0}")
        return count > 0
    except Exception as e:
        logger.error(f"檢查身分證字號是否存在時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        raise Exception(f"檢查身分證字號是否存在時發生錯誤: {str(e)}")

# 驗證身分證字號格式（改進版，提供詳細錯誤訊息）
def validate_id_number(id_number: str) -> Tuple[bool, Optional[str]]:
    if not id_number:
        return False, "請輸入身分證字號"

    if len(id_number) != 10:
        return False, "身分證字號長度必須為 10 個字元"

    if not id_number[0].isalpha() or not id_number[0].isupper():
        return False, "身分證字號第一個字元必須為大寫字母"

    if id_number[1] not in ['1', '2']:
        return False, "身分證字號第二個字元必須為 1 或 2"

    if not id_number[2:].isdigit():
        return False, "身分證字號後 8 個字元必須為數字"

    letter_map = {
        'A': 10, 'B': 11, 'C': 12, 'D': 13, 'E': 14, 'F': 15, 'G': 16, 'H': 17,
        'I': 34, 'J': 18, 'K': 19, 'L': 20, 'M': 21, 'N': 22, 'O': 35, 'P': 23,
        'Q': 24, 'R': 25, 'S': 26, 'T': 27, 'U': 28, 'V': 29, 'W': 32, 'X': 30,
        'Y': 31, 'Z': 33
    }

    first_char_value = letter_map.get(id_number[0])
    if first_char_value is None:
        return False, "身分證字號第一個字元無效"

    total = (first_char_value // 10) + (first_char_value % 10) * 9
    weights = [8, 7, 6, 5, 4, 3, 2, 1]
    for i in range(8):
        total += int(id_number[i + 1]) * weights[i]

    check_digit = (10 - (total % 10)) % 10
    if check_digit != int(id_number[9]):
        return False, "身分證字號檢查碼錯誤"

    return True, None

# 驗證手機號碼格式
def validate_phone_number(phone_number: str) -> bool:
    return len(phone_number) == 10 and phone_number.isdigit()

# 驗證電子郵件格式（簡單檢查）
def validate_email(email: str) -> bool:
    return "@" in email and "." in email

# 驗證密碼長度
def validate_password(password: str) -> bool:
    return len(password) >= 8

# 驗證日期格式
def validate_date(date_str: str) -> bool:
    try:
        datetime.strptime(date_str, "%Y-%m-%d")
        return True
    except ValueError:
        return False

# 檢查角色是否有效
def validate_role(role: str) -> bool:
    return role in VALID_ROLES

# 加密密碼
def hash_password(password: str) -> str:
    return pwd_context.hash(password)

# 驗證密碼
def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

# 插入用戶數據到資料庫
def insert_user(db_config: Dict[str, Any], user_data: Dict[str, Any]):
    try:
        conn = connect_to_db(db_config)
        cursor = conn.cursor()
        query = """
        INSERT INTO users (full_name, gender, birth_date, id_number, password, phone_number, email, role, created_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """
        values = (
            user_data['full_name'],
            user_data['gender'],
            user_data['birth_date'],
            user_data['id_number'],
            user_data['password'],
            user_data['phone_number'],
            user_data['email'],
            user_data['role'],
            datetime.now()
        )
        cursor.execute(query, values)
        conn.commit()
        conn.close()
        logger.info(f"成功插入用戶: id_number={user_data['id_number']}")
    except Exception as e:
        logger.error(f"插入用戶時發生錯誤: id_number={user_data['id_number']}, 錯誤: {str(e)}")
        raise Exception(f"插入用戶時發生錯誤: {str(e)}")

# 檢查用戶是否存在並驗證密碼
def authenticate_user(db_config: Dict[str, Any], id_number: str, password: str, role: str) -> Optional[Dict[str, Any]]:
    try:
        conn = connect_to_db(db_config)
        cursor = conn.cursor()
        query = "SELECT id_number, password, role FROM users WHERE id_number = %s"
        cursor.execute(query, (id_number,))
        user = cursor.fetchone()
        conn.close()
        if user and verify_password(password, user[1]) and user[2] == role:
            logger.info(f"用戶驗證成功: id_number={id_number}, role={role}")
            return {"id_number": user[0], "role": user[2]}
        logger.warning(f"用戶驗證失敗: id_number={id_number}, role={role}")
        return None
    except Exception as e:
        logger.error(f"驗證用戶時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        raise Exception(f"驗證用戶時發生錯誤: {str(e)}")

# 檢查電子郵件是否與身分證字號匹配
def verify_id_number_and_email(db_config: Dict[str, Any], id_number: str, email: str) -> bool:
    try:
        conn = connect_to_db(db_config)
        cursor = conn.cursor()
        query = "SELECT COUNT(*) FROM users WHERE id_number = %s AND email = %s"
        cursor.execute(query, (id_number, email))
        count = cursor.fetchone()[0]
        conn.close()
        return count > 0
    except Exception as e:
        logger.error(f"驗證身分證字號和電子郵件時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        raise Exception(f"驗證身分證字號和電子郵件時發生錯誤: {str(e)}")

# 更新密碼
def update_password(db_config: Dict[str, Any], id_number: str, new_password: str):
    try:
        conn = connect_to_db(db_config)
        cursor = conn.cursor()
        query = "UPDATE users SET password = %s WHERE id_number = %s"
        cursor.execute(query, (hash_password(new_password), id_number))
        conn.commit()
        conn.close()
        logger.info(f"成功更新密碼: id_number={id_number}")
    except Exception as e:
        logger.error(f"更新密碼時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        raise Exception(f"更新密碼時發生錯誤: {str(e)}")

# 提取 PDF 文件的文本
def extract_text_from_pdf(file: BytesIO) -> str:
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        logger.error(f"提取 PDF 文本時發生錯誤: {str(e)}")
        raise Exception(f"提取 PDF 文本時發生錯誤: {str(e)}")

# 提取 Word 文件的文本
def extract_text_from_docx(file: BytesIO) -> str:
    try:
        doc = DocxDocument(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        logger.error(f"提取 Word 文本時發生錯誤: {str(e)}")
        raise Exception(f"提取 Word 文本時發生錯誤: {str(e)}")

# 儲存健康檢查數據到資料庫
def store_health_check(db_config: Dict[str, Any], id_number: str, check_date: str, extracted_text: str):
    try:
        conn = connect_to_db(db_config)
        cursor = conn.cursor()
        query = """
        INSERT INTO health_checks (id_number, check_date, extracted_text, upload_timestamp, data)
        VALUES (%s, %s, %s, %s, %s)
        """
        data_value = json.dumps({"extracted_text": extracted_text})
        cursor.execute(query, (id_number, check_date, extracted_text, datetime.now(), data_value))
        conn.commit()
        conn.close()
        logger.info(f"成功儲存健康檢查數據: id_number={id_number}, check_date={check_date}")
    except Exception as e:
        logger.error(f"儲存健康檢查數據時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        raise Exception(f"儲存健康檢查數據時發生錯誤: {str(e)}")

# 檢索健康檢查數據
def retrieve_health_data(db_config: Dict[str, Any], id_number: str, start_date: Optional[str] = None, end_date: Optional[str] = None) -> List[Dict[str, Any]]:
    try:
        conn = connect_to_db(db_config)
        cursor = conn.cursor()
        query = "SELECT check_date, extracted_text FROM health_checks WHERE id_number = %s"
        params = [id_number]
        if start_date and end_date:
            query += " AND check_date BETWEEN %s AND %s"
            params.extend([start_date, end_date])
        elif start_date:
            query += " AND check_date >= %s"
            params.append(start_date)
        elif end_date:
            query += " AND check_date <= %s"
            params.append(end_date)
        query += " ORDER BY check_date DESC"
        cursor.execute(query, tuple(params))
        records = cursor.fetchall()
        conn.close()
        health_data = [{"check_date": str(record[0]), "extracted_text": record[1]} for record in records]
        logger.info(f"成功檢索健康檢查數據: id_number={id_number}, 記錄數={len(health_data)}")
        return health_data
    except Exception as e:
        logger.error(f"檢索健康檢查數據時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        raise Exception(f"檢索健康檢查數據時發生錯誤: {str(e)}")

# 初始化 LLM
llm = OllamaLLM(
    model="llama3:8b",
    base_url="http://localhost:11434",
    temperature=0.3,
    system_prompt="你是一位醫療助理，專門為台灣用戶提供服務。請始終使用繁體中文回答，並確保回答專業且易於理解。"
)

# 定義分析提示模板
analysis_prompt_template = PromptTemplate(
    input_variables=["health_data", "retrieved_context"],
    template="""
    你是一位專業的醫療助理，專門為台灣用戶提供服務。請使用繁體中文回答，並確保建議專業且易於理解。即使檢索到的醫療知識或健康數據可能包含英文或其他語言，你必須將回答轉換為繁體中文。

    ### 健康檢查數據
    {health_data}

    ### 相關醫療知識
    {retrieved_context}

    ### 分析與建議
    請分析上述健康檢查數據，並結合醫療知識，提供具體的健康建議。如果數據不足以進行分析，請說明需要哪些額外的資訊。回答必須使用繁體中文。
    """
)

# 定義互動提示模板
interactive_prompt_template = PromptTemplate(
    input_variables=["query", "retrieved_context"],
    template="""
    你是一位專業的醫療助理，專門為台灣用戶提供服務。請使用繁體中文回答，並確保建議專業且易於理解。即使用戶查詢或檢索到的醫療知識可能包含英文或其他語言，你必須將回答轉換為繁體中文。

    ### 用戶查詢
    {query}

    ### 相關醫療知識
    {retrieved_context}

    ### 回答
    請根據用戶的查詢和醫療知識，提供具體的回答。如果無法回答，請說明原因並建議如何獲取更多資訊。回答必須使用繁體中文。
    """
)

# 獲取用戶資訊
def get_user_info(db_config: Dict[str, Any], id_number: str) -> Optional[Dict[str, Any]]:
    try:
        conn = connect_to_db(db_config)
        cursor = conn.cursor()
        query = "SELECT id_number, role FROM users WHERE id_number = %s"
        cursor.execute(query, (id_number,))
        user = cursor.fetchone()
        conn.close()
        if user:
            return {"id_number": user[0], "role": user[1]}
        return None
    except Exception as e:
        logger.error(f"獲取用戶資訊時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        raise Exception(f"獲取用戶資訊時發生錯誤: {str(e)}")

# 註冊用戶
def register_user(db_config: Dict[str, Any], request: RegisterRequest) -> Dict[str, str]:
    is_valid, error_message = validate_id_number(request.id_number)
    if not is_valid:
        raise Exception(f"身分證字號格式不正確: {error_message}")

    if check_id_number_exists(db_config, request.id_number):
        raise Exception("身分證字號已存在")

    if not validate_phone_number(request.phone_number):
        raise Exception("手機號碼格式不正確，必須為 10 位數字")

    if not validate_email(request.email):
        raise Exception("電子郵件格式不正確")

    if not validate_password(request.password):
        raise Exception("密碼長度必須至少 8 個字元")

    if not validate_date(request.birth_date):
        raise Exception("出生日期格式不正確，必須為 YYYY-MM-DD")

    if not validate_role(request.role):
        raise Exception(f"無效的角色，必須為 {', '.join(VALID_ROLES)} 之一")

    user_data = {
        "full_name": request.full_name,
        "gender": request.gender,
        "birth_date": request.birth_date,
        "id_number": request.id_number,
        "password": hash_password(request.password),
        "phone_number": request.phone_number,
        "email": request.email,
        "role": request.role
    }

    insert_user(db_config, user_data)
    return {"message": "用戶註冊成功"}

# 登入用戶
def login_user(db_config: Dict[str, Any], request: LoginRequest) -> Dict[str, str]:
    is_valid, error_message = validate_id_number(request.id_number)
    if not is_valid:
        raise Exception(f"身分證字號格式不正確: {error_message}")

    if not validate_role(request.role):
        raise Exception(f"無效的角色，必須為 {', '.join(VALID_ROLES)} 之一")

    user = authenticate_user(db_config, request.id_number, request.password, request.role)
    if not user:
        raise Exception("身分證字號、密碼或角色錯誤")

    return {"id_number": user["id_number"], "role": user["role"]}

# 忘記密碼
def forgot_password(db_config: Dict[str, Any], request: ForgotPasswordRequest) -> Dict[str, str]:
    is_valid, error_message = validate_id_number(request.id_number)
    if not is_valid:
        raise Exception(f"身分證字號格式不正確: {error_message}")

    if not validate_email(request.email):
        raise Exception("電子郵件格式不正確")

    if not verify_id_number_and_email(db_config, request.id_number, request.email):
        raise Exception("身分證字號與電子郵件不匹配")

    new_password = "newpassword123"
    update_password(db_config, request.id_number, new_password)
    return {"message": "密碼已重置，請檢查您的電子郵件（模擬）"}

# 上傳健康檢查數據
async def upload_health_check(db_config: Dict[str, Any], id_number: str, file: Any, check_date: Optional[str] = None) -> Dict[str, str]:
    content = await file.read()
    file_stream = BytesIO(content)

    if file.filename.endswith(".pdf"):
        extracted_text = extract_text_from_pdf(file_stream)
    elif file.filename.endswith(".docx"):
        extracted_text = extract_text_from_docx(file_stream)
    else:
        raise Exception("不支援的文件格式，僅支援 PDF 和 DOCX")

    if not check_date:
        check_date = datetime.now().strftime("%Y-%m-%d")
    else:
        if not validate_date(check_date):
            raise Exception("檢查日期格式不正確，必須為 YYYY-MM-DD")

    store_health_check(db_config, id_number, check_date, extracted_text)
    return {"message": "健康檢查數據上傳成功"}

# 分析健康數據
def analyze_health_data(db_config: Dict[str, Any], id_number: str, start_date: Optional[str] = None, end_date: Optional[str] = None) -> Tuple[Dict[str, Any], Any, Any, Any]:
    health_data = retrieve_health_data(db_config, id_number, start_date, end_date)
    if not health_data:
        logger.warning(f"未找到健康檢查數據: id_number={id_number}")
        result = {
            "health_data": [],
            "analysis_result": "未找到相關健康檢查數據，無法進行分析。",
            "retrieved_context": ""
        }
        return result, llm, ConversationBufferMemory(), interactive_prompt_template

    health_data_str = ""
    for record in health_data:
        health_data_str += f"檢查日期: {record['check_date']}\n提取的文本: {record['extracted_text']}\n\n"

    try:
        query = f"id_number: {id_number}"
        retrieved_docs = vectorstore.similarity_search(query, k=3)
        retrieved_context = ""
        for doc in retrieved_docs:
            retrieved_context += f"{doc.page_content}\n\n"
    except Exception as e:
        logger.error(f"檢索醫療知識時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        retrieved_context = "無法檢索相關醫療知識。"

    try:
        prompt = analysis_prompt_template.format(health_data=health_data_str, retrieved_context=retrieved_context)
        analysis_result = llm.invoke(prompt)
        logger.info(f"成功分析健康數據: id_number={id_number}")
    except Exception as e:
        logger.error(f"LLM 分析健康數據時發生錯誤: id_number={id_number}, 錯誤: {str(e)}")
        analysis_result = "無法進行健康數據分析，請稍後再試。"

    result = {
        "health_data": health_data,
        "analysis_result": analysis_result,
        "retrieved_context": retrieved_context
    }
    return result, llm, ConversationBufferMemory(), interactive_prompt_template

# 互動模式處理用戶查詢
def interactive_query(db_config: Dict[str, Any], query: str) -> Dict[str, Any]:
    id_number = None
    for word in query.split():
        is_valid, _ = validate_id_number(word)
        if is_valid:
            id_number = word
            break

    health_data_str = ""
    if id_number:
        health_data = retrieve_health_data(db_config, id_number)
        for record in health_data:
            health_data_str += f"檢查日期: {record['check_date']}\n提取的文本: {record['extracted_text']}\n\n"

    try:
        retrieved_docs = vectorstore.similarity_search(query, k=3)
        retrieved_context = ""
        for doc in retrieved_docs:
            retrieved_context += f"{doc.page_content}\n\n"
    except Exception as e:
        logger.error(f"檢索醫療知識時發生錯誤: query={query}, 錯誤: {str(e)}")
        retrieved_context = "無法檢索相關醫療知識。"

    try:
        if health_data_str:
            query_with_data = f"{query}\n\n相關健康數據:\n{health_data_str}"
        else:
            query_with_data = query
        prompt = interactive_prompt_template.format(query=query_with_data, retrieved_context=retrieved_context)
        response = llm.invoke(prompt)
        logger.info(f"成功處理互動查詢: query={query}")
    except Exception as e:
        logger.error(f"LLM 處理互動查詢時發生錯誤: query={query}, 錯誤: {str(e)}")
        response = "無法處理您的查詢，請稍後再試。"

    return {
        "query": query,
        "response": response,
        "retrieved_context": retrieved_context
    }

# 填充向量數據庫（知識庫）
def process_batch(docs, embeddings):
    try:
        # 檢查集合是否存在，若不存在則創建
        collections = vectorstore._client.list_collections()
        collection_exists = any(collection.name == vectorstore._collection.name for collection in collections)
        if not collection_exists:
            vectorstore._client.create_collection(
                name=vectorstore._collection.name,
                embedding_function=vectorstore._embedding_function
            )
        # 添加數據到集合
        existing_ids = set(vectorstore._collection.get()['ids'] or [])
        filtered_docs = []
        filtered_embeddings = []
        filtered_metadatas = []
        for doc, emb, meta in zip([doc.page_content for doc in docs], embeddings, [doc.metadata for doc in docs]):
            doc_id = f"{meta['source']}_{meta['entry_id']}"
            if doc_id not in existing_ids:
                filtered_docs.append(doc)
                filtered_embeddings.append(emb)
                filtered_metadatas.append(meta)
        
        if filtered_docs:
            start_time = time.time()
            vectorstore._collection.add(
                documents=filtered_docs,
                embeddings=filtered_embeddings,
                metadatas=filtered_metadatas,
                ids=[f"{meta['source']}_{meta['entry_id']}" for meta in filtered_metadatas]
            )
            elapsed_time = time.time() - start_time
            logger.info(f"批次寫入耗時: {elapsed_time:.2f} 秒，處理 {len(filtered_docs)} 筆數據")
        return True
    except Exception as e:
        logger.error(f"處理批次時發生錯誤: {str(e)}")
        return False

def populate_vectorstore(db_config: Dict[str, Any], toyhom_data_path: str = None, start_idx: int = 0, end_idx: int = None, skip_health_checks: bool = False):
    # 處理 Toyhom 對話數據
    if toyhom_data_path and os.path.exists(toyhom_data_path):
        logger.info(f"開始檢查 Toyhom 數據檔案: {toyhom_data_path}")
        try:
            with open(toyhom_data_path, 'r', encoding='utf-8') as f:
                dialogue_data = json.load(f)
            if not dialogue_data:
                logger.warning(f"Toyhom 對話數據檔案 {toyhom_data_path} 為空")
            else:
                total_entries = len(dialogue_data)
                segment_size = 100000  # 每次處理 100000 筆
                existing_ids = set(vectorstore._collection.get()['ids'] or [])
                for segment_start in range(0, total_entries, segment_size):
                    segment_end = min(segment_start + segment_size, total_entries)
                    logger.info(f"處理 Toyhom 數據段落: {segment_start} 至 {segment_end-1}")
                    dialogue_data_segment = dialogue_data[segment_start:segment_end]
                    dialogue_docs = []
                    skipped_entries = 0
                    for idx, entry in enumerate(dialogue_data_segment, start=segment_start):
                        doc_id = f"Toyhom_{entry.get('source', 'Unknown')}_{idx}"
                        if doc_id in existing_ids:
                            logger.info(f"條目 {idx} 已存在，跳過")
                            skipped_entries += 1
                            continue
                        question = entry.get("question", "")
                        answer = entry.get("answer", "")
                        source = entry.get("source", "Unknown")
                        if not question or not answer:
                            logger.warning(f"條目 {idx} 中缺少 question 或 answer 字段: {entry}")
                            skipped_entries += 1
                            continue
                        content = f"問題: {question}\n回答: {answer}"
                        doc = Document(
                            page_content=content,
                            metadata={"source": f"Toyhom_{source}", "entry_id": idx}
                        )
                        dialogue_docs.append(doc)

                    if not dialogue_docs:
                        logger.warning(f"Toyhom 數據段落 {segment_start} 至 {segment_end-1} 中沒有有效的對話數據，跳過了 {skipped_entries} 個條目")
                    else:
                        max_batch_size = 500  # 調整批次大小
                        total_batches = math.ceil(len(dialogue_docs) / max_batch_size)
                        logger.info(f"準備將 {len(dialogue_docs)} 筆新 Toyhom 對話數據加入知識庫，分為 {total_batches} 批次")

                        texts = [doc.page_content for doc in dialogue_docs]
                        start_time = time.time()
                        embeddings = st_model.encode(texts, batch_size=max_batch_size, show_progress_bar=True, device=device)
                        elapsed_time = time.time() - start_time
                        logger.info(f"嵌入生成耗時: {elapsed_time:.2f} 秒，處理 {len(dialogue_docs)} 筆數據")

                        for i in tqdm(range(0, len(dialogue_docs), max_batch_size), desc=f"載入 Toyhom Dialogue 進度 ({segment_start}-{segment_end-1})", unit="批次"):
                            batch_docs = dialogue_docs[i:i + max_batch_size]
                            batch_embeddings = embeddings[i:i + max_batch_size]
                            success = process_batch(batch_docs, batch_embeddings)
                            if not success:
                                logger.error("某些批次處理失敗，檢查日誌以獲取詳細資訊")

                        new_toyhom_count = len(dialogue_docs)
                        logger.info(f"成功將 {new_toyhom_count} 筆新 Toyhom 對話數據加入知識庫，跳過了 {skipped_entries} 個條目")

                logger.info(f"Toyhom 對話數據處理完成，共 {total_entries} 筆")
        except Exception as e:
            logger.error(f"將 Toyhom 對話數據加入知識庫時發生錯誤: {str(e)}")
            raise
    else:
        logger.warning(f"Toyhom 數據檔案 {toyhom_data_path} 不存在或路徑無效")

    # 載入 health_checks 數據（可選擇性跳過）
    if not skip_health_checks:
        existing_health_ids = set()
        try:
            logger.info("檢查 Chroma 集合中已有的 health_checks 數據...")
            existing_health_docs = vectorstore.get(where={"source": "health_checks"})
            if existing_health_docs and existing_health_docs.get("ids", []):
                existing_health_ids = set(existing_health_docs["ids"])
                logger.info(f"Chroma 集合已包含 {len(existing_health_ids)} 筆 health_checks 數據")
        except Exception as e:
            logger.error(f"檢查 Chroma 集合 health_checks 數據時發生錯誤: {str(e)}")

        try:
            logger.info("開始從 health_checks 表提取數據...")
            conn = connect_to_db(db_config)
            cursor = conn.cursor()
            cursor.execute("SELECT id, id_number, check_date, extracted_text FROM health_checks")
            records = cursor.fetchall()
            new_documents = []
            for record in records:
                record_id = str(record[0])
                if record_id in existing_health_ids:
                    continue
                doc = Document(
                    page_content=f"id_number: {record[1]}\ncheck_date: {record[2]}\nextracted_text: {record[3]}",
                    metadata={"id_number": record[1], "check_date": str(record[2]), "source": "health_checks"}
                )
                new_documents.append(doc)
            if new_documents:
                logger.info(f"準備將 {len(new_documents)} 筆新的健康檢查數據加入知識庫")
                vectorstore.add_documents(new_documents)
                logger.info(f"成功將 {len(new_documents)} 筆新的健康檢查數據加入知識庫")
            else:
                logger.info("沒有新的健康檢查數據需要載入")
            conn.close()
        except Exception as e:
            logger.error(f"將健康檢查數據加入知識庫時發生錯誤: {str(e)}")
            raise

# 初始化知識庫（在程式啟動時執行）
def main():
    try:
        db_config = {
            'server': os.getenv('DB_SERVER'),
            'port': int(os.getenv('DB_PORT')),
            'user': os.getenv('DB_USER'),
            'password': os.getenv('DB_PASSWORD'),
            'database': os.getenv('DB_DATABASE')
        }

        toyhom_data_path = "D:/gg/WOW/medical_dialogue_traditional.json"

        # 設置 skip_health_checks=True 以跳過 health_checks 數據（臨時措施）
        populate_vectorstore(db_config, toyhom_data_path=toyhom_data_path, skip_health_checks=True)
    except Exception as e:
        logger.error(f"主程式執行時發生錯誤: {str(e)}")
        raise

if __name__ == "__main__":
    main()