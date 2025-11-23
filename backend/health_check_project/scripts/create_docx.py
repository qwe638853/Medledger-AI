from docx import Document

doc = Document()
doc.add_heading('健康檢查報告', 0)

doc.add_paragraph('姓名: 王小明  日期: 2024-01-01')

table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '檢查項目'
hdr_cells[1].text = '數值'

data = [
    ('飯前血糖 (Glu-AC)', '200 mg/dL'),
    ('糖化血色素 (HbA1c)', '8.5 %'),
    ('低密度脂蛋白 (LDL-C)', '900 mg/dL'),
    ('肌酸酐 (CRE)', '6.0 mg/dL'),
    ('三酸甘油酯 (TG)', '150 mg/dL'),
    ('白血球 (WBC)', '16.0 10^3/uL')
]

for item, value in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = value

doc.save('my_test_report.docx')
print("✅ 測試檔案 my_test_report.docx 已建立！")