import grpc
from concurrent import futures
import logging
import json
import re
import os
import socket
import io
from pydantic import BaseModel, ValidationError
from typing import List, Optional, Dict, Set
from ollama import Client
import data_pb2, data_pb2_grpc
import time
import pdfplumber
from docx import Document

# 配置日誌
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# --- 環境配置 ---
OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://localhost:11434")
MODEL_NAME = os.getenv("OLLAMA_MODEL", "qwen2.5:14b")

# --- Pydantic 模型 ---
class DiseaseRisk(BaseModel):
    disease: str
    risk_level: str
    main_factors: List[str]
    advice: str

class HealthAnalysis(BaseModel):
    health_score: int
    risk_level_summary: str
    summary: str
    personal_analysis: str
    disease_risks: List[DiseaseRisk]
    insurance_recommendation: List[str]
    protection_plan: List[str]
    success: bool

class DiseaseRiskForInsurer(BaseModel):
    disease: str
    risk_level: str
    main_factors: List[str]
    advice: str

class InsurerAnalysis(BaseModel):
    risk_level_label: str
    summary: str
    core_recommendation: List[str]
    disease_risk_evaluation: List[DiseaseRiskForInsurer]
    success: bool

class HealthAnalysisServicer(data_pb2_grpc.HealthServiceServicer):
    
    # --- 文件解析功能 ---
    def _extract_text_from_pdf_bytes(self, file_bytes):
        text = ""
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"PDF 解析錯誤: {e}")
            return None
        return text

    def _extract_text_from_docx_bytes(self, file_bytes):
        text = ""
        try:
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.error(f"Word 解析錯誤: {e}")
            return None
        return text

    def _parse_health_data_with_llm(self, raw_text):
        target_fields = [
            "Glu-AC", "HbA1c", "LDL-C", "HDL-C", "TG", 
            "ALT(GPT)", "CRE", "Hb", "WBC"
        ]
        prompt = f"""
你是一個專業的醫療數據錄入員。請從以下的「健檢報告原始文本」中提取數值，並輸出為 JSON 格式。

=== 健檢報告原始文本 ===
{raw_text[:3000]} 
=== 結束 ===

=== 提取規則 ===
1. 請提取以下指標的數值：{', '.join(target_fields)}
2. 如果找不到某個指標，該欄位的值請填 "N/A"。
3. 數值請保留單位（例如 "90 mg/dL" 或 "5.6%"）。
4. **嚴格只輸出 JSON**，不要有任何 Markdown 標記或額外文字。

=== 輸出範例 ===
{{
  "Glu-AC": "95 mg/dL",
  "HbA1c": "5.2%",
  ...
}}
"""
        try:
            response = self.ollama_client.chat(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": prompt}]
            )
            content = response['message']['content'].strip()
            # 清理 Markdown
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
            
            json.loads(content)
            return content
        except Exception as e:
            logger.error(f"LLM 解析失敗: {e}")
            return None

    def ParseDocument(self, request, context):
        logger.info(f"收到文件解析請求，類型: {request.file_type}, 大小: {len(request.file_content)} bytes")
        
        raw_text = ""
        if request.file_type.lower() == "pdf":
            raw_text = self._extract_text_from_pdf_bytes(request.file_content)
        elif request.file_type.lower() == "docx":
            raw_text = self._extract_text_from_docx_bytes(request.file_content)
        else:
            return data_pb2.ParseDocumentResponse(success=False, error_message="不支援的檔案格式")

        if not raw_text:
            return data_pb2.ParseDocumentResponse(success=False, error_message="無法提取文字內容")

        json_result = self._parse_health_data_with_llm(raw_text)
        
        if json_result:
            logger.info("文件解析成功")
            return data_pb2.ParseDocumentResponse(success=True, result_json=json_result)
        else:
            return data_pb2.ParseDocumentResponse(success=False, error_message="AI 解析結構化數據失敗")

    # --- RAG 邏輯 ---
    def _load_health_rules(self) -> (str, Dict[str, str]):
        # [路徑更新] 這裡改成指向 data/ 資料夾
        rules_file = "data/health_rules.txt"
        
        if not os.path.exists(rules_file):
            logger.error(f"關鍵錯誤：知識文件 '{rules_file}' 不存在！")
            return "無可用的醫學參考資料。", {}
            
        try:
            with open(rules_file, 'r', encoding='utf-8') as f:
                full_text = f.read()
            logger.info(f"成功載入 '{rules_file}' (大小: {len(full_text)} 字節)")
            
            chunks = re.split(r'(### .*\n)', full_text)
            knowledge_chunks: Dict[str, str] = {}
            
            if not chunks[0].startswith("### "):
                knowledge_chunks["General_Header"] = chunks[0].strip()
                chunks = chunks[1:]

            for i in range(0, len(chunks), 2):
                if i + 1 < len(chunks):
                    header = chunks[i].strip()
                    content = chunks[i+1].strip()
                    knowledge_chunks[header] = content
                
            logger.info(f"(RAG 4.1) 知識庫已切割為 {len(knowledge_chunks)} 個主題塊。")
            return full_text, knowledge_chunks
            
        except Exception as e:
            logger.error(f"載入 '{rules_file}' 失敗: {str(e)}")
            return "讀取醫學參考資料失敗。", {}

    def _get_key_to_topic_map(self) -> Dict[str, List[str]]:
        return {
            "血糖指標": ["Glu-AC", "HbA1c", "Glu-PC"],
            "肝功能指標": ["Alb", "TP", "AST", "ALT", "D-Bil", "ALP", "T-Bil"],
            "腎功能指標": ["UN", "CRE", "U.A", "Alb/CRE Ratio"],
            "血脂指標": ["T-CHO", "LDL-C", "HDL-C", "TG"],
            "全血球計數": ["Hb", "Hct", "PLT", "WBC", "RBC", "MCV", "MCH", "MCHC", "RDW-CV"],
            "凝血功能": ["PT", "aPTT"],
            "發炎指標": ["hsCRP", "ESR"],
            "癌症指標": ["AFP", "CEA", "CA-125", "CA19-9"],
            "尿液常規檢查": ["Specific Gravity", "PH", "Protein", "Glucose", "Bilirubin", "Urobilinogen", "RBC", "WBC", "Epithelial", "Casts", "Ketone", "Crystal", "Bacteria", "Albumin", "Creatinine", "Nitrite", "Occult Blood", "WBC Esterase"],
            "其他": ["BP"],
            "核心評分準則": ["*ALWAYS_INCLUDE*"]
        }

    def _filter_context_for_rag(self, test_results: dict) -> str:
        input_keys_raw = set(test_results.keys())
        relevant_topics: Set[str] = set() 
        topic_map = self._get_key_to_topic_map()
        
        for topic, indicators in topic_map.items():
            if "*ALWAYS_INCLUDE*" in indicators:
                relevant_topics.add(topic)
                continue
            for key in input_keys_raw:
                clean_key = key.split('(')[0].strip() 
                if clean_key in indicators:
                    relevant_topics.add(topic)
                    break 

        logger.info(f"(RAG 4.1) 偵測到相關主題: {relevant_topics}")

        context_parts: List[str] = []
        for header, content in self.knowledge_chunks.items():
            for topic in relevant_topics:
                if topic in header: 
                    context_parts.append(f"{header}\n{content}")
                    break

        if not context_parts:
            logger.warning("(RAG 4.1) 警告：未篩選到任何相關規則！")
            return "無相關醫學參考資料"
            
        final_context = "\n---\n".join(context_parts)
        logger.info(f"(RAG 4.1) 篩選後的上下文 (大小: {len(final_context)} 字節)")
        return final_context

    def __init__(self):
        try:
            logger.info("開始初始化服務器依賴")
            self.full_health_rules, self.knowledge_chunks = self._load_health_rules()
            
            if not self._check_ollama_connection():
                raise RuntimeError(f"Ollama 服務不可用: {OLLAMA_HOST}")
                
            self.ollama_client = Client(host=OLLAMA_HOST)
            logger.info(f"Ollama 客戶端初始化，目標地址: {OLLAMA_HOST}")
            
        except Exception as e:
            logger.error(f"服務器初始化失敗: {str(e)}")
            raise

    def _check_ollama_connection(self, max_retries=3, delay=2) -> bool:
        try:
            host_port = OLLAMA_HOST.replace("http://", "").replace("https://", "")
            if ":" in host_port: host, port = host_port.split(":")
            else: host, port = host_port, 11434
            logger.info(f"檢查 Ollama 連接: {host}:{port}")
            for attempt in range(max_retries):
                with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                    s.settimeout(5)
                    result = s.connect_ex((host, int(port)))
                    if result == 0:
                        return True
                    time.sleep(delay)
            logger.error("Ollama 連線檢查失敗，所有重試均失敗")
            return False
        except Exception as e:
            logger.error(f"Ollama 連線檢查失敗: {str(e)}")
            return False

    def _call_ollama_json(self, prompt: str, schema: BaseModel) -> Optional[BaseModel]:
        logger.info(f"發送 Ollama prompt...")
        try:
            response = self.ollama_client.chat(
                model=MODEL_NAME,
                messages=[{"role": "user", "content": prompt}],
                format=schema.model_json_schema()
            )
            content = response['message']['content']
            return schema.model_validate_json(content)
        except Exception as e:
            logger.error(f"Ollama 錯誤: {str(e)}")
            return None

    def AnalyzeHealthReportForUser(self, request, context):
        logger.info(f"處理用戶健康報告，報告 ID: {request.report_id}")
        try:
            test_results = json.loads(request.test_results_json)
            available_keys = list(test_results.keys())
            query_text = "\n".join([f"{k}: {v}" for k, v in test_results.items()])
            
            full_context = self._filter_context_for_rag(test_results)

            prompt = f"""
你是一位非常親切、有耐心的台灣家庭醫師。你的溝通對象是**完全沒有醫學背景的一般民眾**。
請使用最**白話、通俗、易懂**的繁體中文來解釋健康檢查報告。

=== 健康檢查數據 ===
{query_text}

=== 相關醫學知識 (專業參考) ===
{full_context}

=== 解說原則 (嚴格遵守) ===
1. **拒絕術語**：如果一定要用到醫學名詞，必須立刻用白話解釋。
2. **善用比喻**：請多用生活中的例子來比喻。
3. **語氣溫暖**：不要冷冰冰的，要像朋友一樣關心。
4. **重點優先**：先講結論，再講為什麼。
5. **不打分數**：只用文字描述風險等級。

=== 輸出格式（JSON）===
{{
  "risk_level_summary": "<總體健康狀態短語>",
  "summary": "<白話文健康總結>",
  "personal_analysis": "<詳細白話解說>",
  "disease_risks": [
    {{
      "disease": "<疾病名稱>", 
      "risk_level": "<風險等級>", 
      "main_factors": ["<異常指標>"], 
      "advice": "<白話建議>"
    }}
  ],
  "insurance_recommendation": ["<白話保險建議>"],
  "protection_plan": ["<白話健康計畫>"],
  "success": true
}}
"""
            result = self._call_ollama_json(prompt, HealthAnalysis)
            if not result: raise ValueError("LLM 回傳無效格式")

            disease_risks_proto = [
                data_pb2.DiseaseRisk(
                    disease=dr.disease, risk_level=dr.risk_level,
                    main_factors=dr.main_factors, advice=dr.advice
                ) for dr in result.disease_risks
            ]
            
            final_summary = f"【{result.risk_level_summary}】\n{result.summary}"

            return data_pb2.UserHealthAnalysisResponse(
                health_score=0, summary=final_summary, personal_analysis=result.personal_analysis,
                disease_risks=disease_risks_proto,
                insurance_recommendation=result.insurance_recommendation,
                protection_plan=result.protection_plan,
                success=result.success
            )

        except Exception as e:
            logger.error(f"用戶分析失敗: {str(e)}")
            return data_pb2.UserHealthAnalysisResponse(
                health_score=0, summary="分析失敗", personal_analysis="請稍後重試", success=False
            )

    def AnalyzeHealthReportForInsurer(self, request, context):
        logger.info(f"處理保險公司健康報告，報告 ID: {request.report_id}")
        disease_risk_proto_list = []
        try:
            test_results = json.loads(request.test_results_json)
            available_keys = list(test_results.keys())
            query_text = "\n".join([f"{k}: {v}" for k, v in test_results.items()])
            full_context = self._filter_context_for_rag(test_results)

            prompt = f"""
你是一位專業的保險核保風險顧問。請使用**繁體中文**為保險公司分析客戶的健康風險。

=== 健康檢查數據 ===
{query_text}

=== 相關醫學知識 ===
{full_context}

=== 核心任務 ===
1. **客觀分析**：請向核保人員解釋這些異常數據對「未來理賠機率」的影響。
2. **白話解釋風險**：即便對象是核保員，也請清楚說明病理機制導致的長期風險。
3. **不做決定**：提供風險評估，但不直接下達拒保指令。
4. **不打分數**：只提供風險等級。

=== 輸出格式（JSON）===
{{
  "risk_level_label": "<總體風險等級>",
  "summary": "<詳細風險摘要>",
  "core_recommendation": ["<風險管理建議>"],
  "disease_risk_evaluation": [
    {{
      "disease": "<疾病名稱>", 
      "risk_level": "<風險等級>", 
      "main_factors": ["<異常指標>"], 
      "advice": "<詳細風險分析>"
    }}
  ],
  "success": true
}}
"""
            result = self._call_ollama_json(prompt, InsurerAnalysis)
            if not result: raise ValueError("LLM 回傳無效格式")

            disease_risk_proto_list = [
                data_pb2.DiseaseRiskForInsurer(
                    disease=dr.disease, risk_score=0, risk_level=dr.risk_level,
                    main_factors=dr.main_factors, advice=dr.advice
                ) for dr in result.disease_risk_evaluation
            ]
            return data_pb2.InsurerHealthAnalysisResponse(
                overall_risk_score=0, risk_level_label=result.risk_level_label,
                summary=result.summary, core_recommendation=result.core_recommendation,
                disease_risk_evaluation=disease_risk_proto_list, success=result.success
            )
        except Exception as e:
            logger.error(f"保險分析失敗: {str(e)}")
            return data_pb2.InsurerHealthAnalysisResponse(
                overall_risk_score=0, risk_level_label="分析失敗", summary=f"分析失敗: {str(e)}",
                success=False
            )

def serve():
    try:
        grpc_port = os.getenv("PYTHON_BACKEND_GRPC_PORT", "50052")
        grpc_host = os.getenv("PYTHON_BACKEND_GRPC_HOST", "[::]")
        server = grpc.server(futures.ThreadPoolExecutor(max_workers=5))
        data_pb2_grpc.add_HealthServiceServicer_to_server(HealthAnalysisServicer(), server)
        listen_addr = f"{grpc_host}:{grpc_port}"
        server.add_insecure_port(listen_addr)
        logger.info(f"Python gRPC 服務器啟動，監聽地址: {listen_addr}")
        logger.info(f"注意: Go Server 運行在 :50051，Python Backend 運行在 :{grpc_port}")
        server.start()
        try:
            while True: server.wait_for_termination(timeout=60)
        except KeyboardInterrupt:
            server.stop(grace=None)
    except Exception as e:
        logger.error(f"服務器啟動失敗: {str(e)}")
        raise

if __name__ == "__main__":
    serve()